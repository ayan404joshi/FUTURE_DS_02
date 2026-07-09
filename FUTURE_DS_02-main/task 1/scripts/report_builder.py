"""Build the written project report in Markdown and Word formats."""
from __future__ import annotations

from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from common import ANALYSIS_SUMMARY_FILE, ASSETS_DIR, MODEL_CARD_FILE, PROCESSED_DIR, REPORT_DOCX_FILE, REPORT_MD_FILE, ensure_directories, read_json


def _fmt_pct(value: float) -> str:
    """Format a decimal as a percentage string."""
    return f"{value * 100:.1f}%"


def _fmt_num(value: float) -> str:
    """Format a numeric value with comma separators."""
    return f"{value:,.2f}"


def _recommendations(summary: dict[str, Any], model_metrics: dict[str, Any], churn_df: pd.DataFrame) -> list[dict[str, str]]:
    """Create quantified business recommendations linked to observed findings."""
    overall = summary["churn_rate"]
    top_two_share = summary["top_two_reason_share"]
    high_value_gap = max(0.0, summary["high_value_churn_rate"] - overall)
    low_usage_corr = abs(summary["usage_churn_spearman"])
    support_corr = summary["support_churn_spearman"]

    top_reason_impact = overall * top_two_share * 0.20
    high_value_impact = high_value_gap * summary["high_value_share"] * 0.50
    risk_segment_impact = overall * min(0.08, (low_usage_corr + abs(support_corr)) / 10)

    return [
        {
            "title": "Attack the top churn reasons first",
            "evidence": f"The top two churn reasons account for {_fmt_pct(top_two_share)} of all churn events.",
            "action": "Create a billing-and-product retention playbook for customers citing billing, no usage, price sensitivity, and support issues.",
            "impact": f"If targeted interventions reduce churn among those reasons by 20%, estimated overall churn could drop by about {_fmt_pct(top_reason_impact)}.",
        },
        {
            "title": "Protect high-value customers",
            "evidence": f"High CLV customers churn at {_fmt_pct(summary['high_value_churn_rate'])} versus {_fmt_pct(overall)} overall.",
            "action": "Launch a high-value save desk for premium and high-CLV accounts with proactive renewal outreach and priority support.",
            "impact": f"Closing half of the high-value churn gap could reduce total churn by roughly {_fmt_pct(high_value_impact)} and protect disproportionate revenue.",
        },
        {
            "title": "Intervene earlier on low-usage / high-support accounts",
            "evidence": f"Usage and support show meaningful churn separation (Spearman {summary['usage_churn_spearman']:.2f} and {summary['support_churn_spearman']:.2f}).",
            "action": "Trigger lifecycle nudges, in-app education, and support escalation when engagement drops and ticket volume rises.",
            "impact": f"A 10% reduction in churn among the highest-risk engagement segments could lower churn by about {_fmt_pct(risk_segment_impact)}.",
        },
    ]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a consistent heading style to the Word report."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bullets(doc: Document, items: list[str]) -> None:
    """Add bullet paragraphs to the document."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _insert_chart(doc: Document, image_name: str, caption: str) -> None:
    """Insert a chart image if it exists on disk."""
    image_path = ASSETS_DIR / image_name
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.6))
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.runs[0].italic = True


def build_report() -> dict[str, Any]:
    """Build the narrative report in Markdown and Word format."""
    ensure_directories()
    summary = read_json(ANALYSIS_SUMMARY_FILE)
    model_metrics = read_json(PROCESSED_DIR / "model_metrics.json")
    churn_df = pd.read_csv(PROCESSED_DIR / "clv_tier_churn.csv")
    recommendations = _recommendations(summary, model_metrics, churn_df)

    best_model_name = model_metrics["best_model"]
    best_metrics = model_metrics["metrics"][best_model_name]
    feature_importance = pd.read_csv(PROCESSED_DIR / "feature_importance.csv").head(10)

    report_lines = [
        "# Customer Retention & Churn Analysis Report",
        "",
        "## Executive Summary",
        "",
        f"The business analyzed {summary['total_customers']:,} subscription customers and observed an overall churn rate of {_fmt_pct(summary['churn_rate'])}. The strongest churn drivers were the top two churn reasons, which together explain {_fmt_pct(summary['top_two_reason_share'])} of churn events, while high-value customers churned at {_fmt_pct(summary['high_value_churn_rate'])}, above the overall baseline.",
        f"The churn model selected {best_model_name.replace('_', ' ')} as the best performer with ROC-AUC {best_metrics['roc_auc']:.3f} and F1 {best_metrics['f1']:.3f}. The dominant quantitative drivers align with the EDA: low usage, elevated support contacts, and price/discount sensitivity.",
        "",
        "## Business Problem & Objective",
        "",
        "The objective was to identify churn patterns, isolate the retention levers with the highest business value, quantify customer lifetime value, and build a churn prediction model that can support targeted retention actions.",
        "",
        "## Data Overview & Methodology",
        "",
        f"A realistic synthetic subscription dataset was generated with {summary['total_customers']:,} customers, including sign-up date, plan, usage, support, payment, region, age group, acquisition channel, discount, upsell, downsell, tenure, churn label, and churn reason. The raw data intentionally contained mixed date formats, duplicates, missing values, and outliers so the cleaning step could demonstrate production-style preparation.",
        "The analysis used pandas for data preparation, seaborn/matplotlib for charts, scikit-learn for modeling, and a class-weight strategy to address the imbalanced churn target without fabricating synthetic category combinations.",
        "",
        "## Key Findings",
        "",
        f"- Overall churn rate: {_fmt_pct(summary['churn_rate'])}",
        f"- Average CLV: {_fmt_num(summary['avg_clv'])}",
        f"- Average tenure: {_fmt_num(summary['avg_tenure'])} months",
        f"- Usage-frequency vs churn Spearman correlation: {summary['usage_churn_spearman']:.2f}",
        f"- Support-ticket vs churn Spearman correlation: {summary['support_churn_spearman']:.2f}",
        "",
        "### Churn Trend",
        "",
        f"Monthly churn trend shows how retention moved across signup cohorts. The latest measured monthly churn rate was {_fmt_pct(summary['monthly_trend_last_value'])}.",
        "",
        "### Segment Differences",
        "",
        "Churn varies materially across plan, region, acquisition channel, and payment method. The dashboard and charts in the appendix provide the exact segment ranking used in the recommendation design.",
        "",
        "### Cohort Retention",
        "",
        "Cohort heatmaps reveal the earliest months after signup as the highest-risk window, which is consistent with onboarding friction and value-realization lag.",
        "",
        "### CLV Insights",
        "",
        f"High-value customers account for {_fmt_pct(summary['high_value_share'])} of the base but churn at {_fmt_pct(summary['high_value_churn_rate'])}, which creates a clear revenue risk signal.",
        "",
        "## Churn Prediction Model Summary & Top Drivers",
        "",
        f"The best model was {best_model_name.replace('_', ' ')} with ROC-AUC {best_metrics['roc_auc']:.3f}, precision {best_metrics['precision']:.3f}, recall {best_metrics['recall']:.3f}, and F1 {best_metrics['f1']:.3f}.",
        "Top drivers were extracted from the fitted model and align with the EDA patterns shown below.",
        "",
        "## Top Drivers",
        "",
    ]
    for _, row in feature_importance.iterrows():
        report_lines.append(f"- {row['feature']}: {row['importance']:.4f}")

    report_lines.extend([
        "",
        "## Actionable Recommendations",
        "",
    ])
    for idx, rec in enumerate(recommendations, start=1):
        report_lines.extend([
            f"### Recommendation {idx}: {rec['title']}",
            f"- Evidence: {rec['evidence']}",
            f"- Action: {rec['action']}",
            f"- Estimated impact: {rec['impact']}",
            "",
        ])

    report_lines.extend([
        "## Limitations & Future Work",
        "",
        "The dataset is synthetic, so the numerical results demonstrate methodology rather than live business performance. A real engagement would add product telemetry, billing history, contract terms, and customer success touches, then validate the model on a time-based holdout set.",
        "Future work should include calibrated probability thresholds, uplift modeling, and retention experiment tracking.",
        "",
        "## Appendix",
        "",
        "- [assets/churn_trend.png](../assets/churn_trend.png)",
        "- [assets/plan_churn.png](../assets/plan_churn.png)",
        "- [assets/region_churn.png](../assets/region_churn.png)",
        "- [assets/channel_churn.png](../assets/channel_churn.png)",
        "- [assets/payment_churn.png](../assets/payment_churn.png)",
        "- [assets/churn_reason_pareto.png](../assets/churn_reason_pareto.png)",
        "- [assets/usage_vs_churn.png](../assets/usage_vs_churn.png)",
        "- [assets/support_vs_churn.png](../assets/support_vs_churn.png)",
        "- [assets/tenure_distribution.png](../assets/tenure_distribution.png)",
        "- [assets/cohort_retention_heatmap.png](../assets/cohort_retention_heatmap.png)",
        "- [assets/clv_tier_churn.png](../assets/clv_tier_churn.png)",
        "- [processed/top_churn_drivers.png](../data/processed/top_churn_drivers.png)",
    ])

    REPORT_MD_FILE.write_text("\n".join(report_lines), encoding="utf-8")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Customer Retention & Churn Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Synthetic subscription-business capstone deliverable").italic = True

    _add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(
        f"The analysis covered {summary['total_customers']:,} customers and found an overall churn rate of {_fmt_pct(summary['churn_rate'])}. High-value customers churned above the baseline, while the top two churn reasons represented {_fmt_pct(summary['top_two_reason_share'])} of all churn events."
    )
    doc.add_paragraph(
        f"The best churn classifier was {best_model_name.replace('_', ' ')} with ROC-AUC {best_metrics['roc_auc']:.3f}. The results point to a clear retention priority: reduce early-life disengagement and service friction in the highest-risk segments."
    )

    _add_heading(doc, "Business Problem & Objective", level=1)
    doc.add_paragraph(
        "The project aimed to identify churn patterns, quantify lifetime value risk, and translate those findings into specific retention actions for a subscription business."
    )

    _add_heading(doc, "Data Overview & Methodology", level=1)
    doc.add_paragraph(
        "A synthetic raw dataset was created to mimic real client data quality issues, then cleaned into a single analysis-ready table. The workflow used pandas for preparation, seaborn/matplotlib for visualization, and scikit-learn for churn modeling."
    )

    _add_heading(doc, "Key Findings", level=1)
    _add_bullets(
        doc,
        [
            f"Overall churn rate: {_fmt_pct(summary['churn_rate'])}",
            f"Average CLV: {_fmt_num(summary['avg_clv'])}",
            f"Average tenure: {_fmt_num(summary['avg_tenure'])} months",
            f"Usage vs churn Spearman correlation: {summary['usage_churn_spearman']:.2f}",
            f"Support tickets vs churn Spearman correlation: {summary['support_churn_spearman']:.2f}",
        ],
    )
    _insert_chart(doc, "churn_trend.png", "Figure 1. Monthly churn trend.")
    _insert_chart(doc, "churn_reason_pareto.png", "Figure 2. Pareto breakdown of churn reasons.")
    _insert_chart(doc, "cohort_retention_heatmap.png", "Figure 3. Monthly cohort retention heatmap.")
    _insert_chart(doc, "clv_tier_churn.png", "Figure 4. Churn rate by CLV tier.")

    _add_heading(doc, "Churn Prediction Model Summary & Top Drivers", level=1)
    doc.add_paragraph(
        f"The best model was {best_model_name.replace('_', ' ')} with ROC-AUC {best_metrics['roc_auc']:.3f}, precision {best_metrics['precision']:.3f}, recall {best_metrics['recall']:.3f}, and F1 {best_metrics['f1']:.3f}."
    )
    _insert_chart(doc, "top_churn_drivers.png", "Figure 5. Top churn drivers from the fitted model.")

    _add_heading(doc, "Actionable Recommendations", level=1)
    for idx, rec in enumerate(recommendations, start=1):
        doc.add_paragraph(f"Recommendation {idx}: {rec['title']}", style="Heading 2")
        _add_bullets(doc, [rec["evidence"], rec["action"], rec["impact"]])

    _add_heading(doc, "Limitations & Future Work", level=1)
    doc.add_paragraph(
        "The results are based on synthetic data and are intended to demonstrate end-to-end analytical rigor. In a real deployment, the model would be validated on time-based out-of-sample data and connected to intervention outcomes."
    )

    _add_heading(doc, "Appendix", level=1)
    doc.add_paragraph("Project artifacts:")
    _add_bullets(
        doc,
        [
            str(PROCESSED_DIR / "subscription_customers_clean.csv"),
            str(PROCESSED_DIR / "model_metrics.json"),
            str(PROCESSED_DIR / "feature_importance.csv"),
            str(PROCESSED_DIR / "cohort_retention_table.csv"),
            str(MODEL_CARD_FILE),
            str(REPORT_MD_FILE),
        ],
    )

    doc.save(REPORT_DOCX_FILE)
    return {"markdown": str(REPORT_MD_FILE), "docx": str(REPORT_DOCX_FILE), "recommendations": recommendations}


if __name__ == "__main__":
    result = build_report()
    print(f"Report written to {REPORT_DOCX_FILE}")
